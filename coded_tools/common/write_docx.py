"""
coded_tools/common/write_docx.py
â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
WriteDocx â€” create or append to a Cognizant-branded Word (.docx) file.

Uses Cognizant_Template.docx by default (coded_tools/assets/).
Falls back to a plain blank document if the template is not found.

Parses the content string and converts it to a formatted Word document:
  â€¢ "# Title"      â†’ Heading 1
  â€¢ "## Section"   â†’ Heading 2
  â€¢ "### Sub"      â†’ Heading 3
  â€¢ "- item"       â†’ Bullet list item
  â€¢ "|col1|col2|"  â†’ Table row (first row becomes bold header)
  â€¢ "---"          â†’ Page break
  â€¢ Anything else  â†’ Normal paragraph

Modes:
  â€¢ mode="write"  (default) â€” create a fresh document from the Cognizant template
  â€¢ mode="append"           â€” open an existing .docx and add content at the end;
                              if the file does not exist yet, creates it fresh

HOCON class reference
â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    "class": "coded_tools.common.write_docx.WriteDocx"

HOCON tool block (copy-paste into any network)
â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    {
        "name": "write_docx",
        "class": "coded_tools.common.write_docx.WriteDocx",
        "function": {
            "description": "Create or append to a Cognizant-branded Word (.docx) file. Use # for headings, - for bullets, |col|col| for tables, --- for page breaks. Use mode='append' to add content to an existing file.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path":          { "type": "string",  "description": "Relative path for the output file inside the output directory (e.g. 'outputs/report.docx')." },
                    "content":       { "type": "string",  "description": "Document content using markdown-like syntax: # H1, ## H2, ### H3, - bullet, |col|col| table rows, --- page break, plain text for paragraphs." },
                    "title":         { "type": "string",  "description": "Optional document title added as a cover heading (write mode only)." },
                    "mode":          { "type": "string",  "description": "'write' (default â€” create fresh from Cognizant template) or 'append' (add content to existing file; creates fresh if file does not exist)." },
                    "template_path": { "type": "string",  "description": "Optional path to a custom .docx template. Overrides the default Cognizant template (write mode only)." },
                    "agent":         { "type": "string",  "description": "Calling agent name (audit log)." }
                },
                "required": ["path", "content"]
            }
        }
    }

sly_data keys read
â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    output_dir     (preferred) â€” where relative output paths are resolved
    workspace_dir  (fallback)
    project_folder (fallback)  â€” bidmagic/dealcraft compat

Template
â”€â”€â”€â”€â”€â”€â”€â”€â”€
    Default: coded_tools/assets/Cognizant_Template.docx
    Fallback: blank document (if template file is missing)
    Ignored in mode="append" â€” the existing file's styles are preserved
"""

from __future__ import annotations

import logging
import os
from typing import Any

from neuro_san.interfaces.coded_tool import CodedTool

from coded_tools.common._base import get_output_dir, log_call, resolve_output_path

logger = logging.getLogger(__name__)

# Template lives in coded_tools/assets/ â€” one level up from coded_tools/common/
_ASSETS_DIR    = os.path.join(os.path.dirname(__file__), "..", "assets")
_DOCX_TEMPLATE = os.path.abspath(os.path.join(_ASSETS_DIR, "Cognizant_Template.docx"))


def _load_document(template_override: str | None, sly_data: dict):
    """Load the document from template or blank. Returns (Document, used_template: bool)."""
    from docx import Document  # type: ignore

    # Custom template override from args
    if template_override:
        tpath = template_override if os.path.isabs(template_override) \
            else os.path.join(get_output_dir(sly_data), template_override)
        if os.path.exists(tpath):
            return Document(tpath), True
        logger.warning("WriteDocx: custom template not found at '%s' â€” falling back", tpath)

    # Default Cognizant template
    if os.path.exists(_DOCX_TEMPLATE):
        return Document(_DOCX_TEMPLATE), True

    logger.warning("WriteDocx: Cognizant_Template.docx not found at '%s' â€” using blank doc", _DOCX_TEMPLATE)
    return Document(), False


def _clear_template_content(doc):
    """Remove all paragraphs from a freshly loaded template (keeps styles intact)."""
    for para in list(doc.paragraphs):
        p = para._element
        p.getparent().remove(p)


def _parse_and_build(doc, content: str, title: str | None):
    """Parse content lines and add them to the Word document."""
    from docx.oxml.ns import qn        # type: ignore
    from docx.oxml import OxmlElement  # type: ignore

    if title:
        doc.add_heading(title, level=0)

    lines = content.splitlines()
    i = 0
    table_rows: list[list[str]] = []

    def _flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        cols = max(len(r) for r in table_rows)
        t = doc.add_table(rows=0, cols=cols)
        t.style = "Table Grid"
        for r_idx, cells in enumerate(table_rows):
            row = t.add_row()
            for c_idx, text in enumerate(cells[:cols]):
                cell = row.cells[c_idx]
                cell.text = text
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Table row
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped[1:-1].split("|")]
            # Skip markdown separator rows like |---|---|
            if not all(set(c) <= {"-", ":"} for c in cells):
                table_rows.append(cells)
            i += 1
            continue
        else:
            _flush_table()

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped == "---":
            para = doc.add_paragraph()
            run = para.add_run()
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            run._r.append(br)
        elif stripped == "":
            pass  # blank lines skipped â€” paragraphs auto-spaced
        else:
            doc.add_paragraph(stripped)

        i += 1

    _flush_table()


class WriteDocx(CodedTool):
    """Create a Cognizant-branded Word (.docx) file from markdown-like structured text."""

    async def async_invoke(self, args: dict[str, Any], sly_data: dict[str, Any]) -> str:
        try:
            from docx import Document  # type: ignore  # noqa: F401 â€” verify import
        except ImportError:
            return "Error: python-docx is not installed. Run: pip install python-docx"

        path_raw          = (args.get("path")          or "").strip()
        content           =  args.get("content")
        title             = (args.get("title")         or "").strip() or None
        mode_raw          = (args.get("mode")          or "write").strip().lower()
        template_override = (args.get("template_path") or "").strip() or None
        agent             = (args.get("agent")         or "unknown-agent").strip()

        if not path_raw:
            return "Error: write_docx requires 'path'."
        if content is None:
            return "Error: write_docx requires 'content'."
        if mode_raw not in ("write", "append"):
            return f"Error: 'mode' must be 'write' or 'append' (got {mode_raw!r})."
        if not path_raw.lower().endswith(".docx"):
            path_raw = path_raw + ".docx"

        try:
            abs_path = resolve_output_path(path_raw, sly_data)
        except ValueError as exc:
            log_call(sly_data, tool="WriteDocx", agent=agent, target=path_raw,
                     status="ERROR", detail=str(exc))
            return f"Error: {exc}"

        try:
            os.makedirs(os.path.dirname(abs_path), exist_ok=True)

            if mode_raw == "append" and os.path.exists(abs_path):
                # Open existing document and add content at the end
                from docx import Document  # type: ignore
                doc = Document(abs_path)
                _parse_and_build(doc, content, title=None)   # title ignored in append
                mode_note = "append"
                used_template = False
            else:
                # Create fresh from template (write mode, or append on missing file)
                doc, used_template = _load_document(template_override, sly_data)
                if used_template:
                    _clear_template_content(doc)
                _parse_and_build(doc, content, title)
                mode_note = "write" if mode_raw == "write" else "append-as-create"

            doc.save(abs_path)
            file_size = os.path.getsize(abs_path)

        except Exception as exc:
            log_call(sly_data, tool="WriteDocx", agent=agent, target=path_raw,
                     status="ERROR", detail=str(exc))
            logger.exception("WriteDocx failed for %s", path_raw)
            return f"Error: failed to write '{path_raw}': {exc}"

        if mode_note == "append":
            template_note = "appended to existing file"
        elif used_template:
            template_note = "Cognizant template"
        else:
            template_note = "plain (template not found)"

        log_call(sly_data, tool="WriteDocx", agent=agent, target=path_raw,
                 status="OK", detail=f"{mode_note}: {file_size} bytes ({template_note})")
        logger.debug("WriteDocx: %s (%d bytes, %s, %s)", path_raw, file_size, mode_note, template_note)

        verb = "appended to" if mode_note == "append" else "wrote"
        return f"OK: {verb} '{path_raw}' ({file_size} bytes, {template_note})."

